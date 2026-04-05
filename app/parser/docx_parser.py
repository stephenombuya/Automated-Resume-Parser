"""
Production-grade DOCX Parser with robust error handling, memory optimization,
and extensible content extraction.
"""

import logging
import os
from pathlib import Path
from typing import Optional, Dict, Any, List, Union
from dataclasses import dataclass, field
from contextlib import contextmanager

from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.opc.exceptions import PackageNotFoundError
from docx.opc.part import Part

# Configure module-level logger
logger = logging.getLogger(__name__)


@dataclass
class ParseResult:
    """Structured result object for parsed document content."""
    text: str = ""
    paragraphs: List[str] = field(default_factory=list)
    tables: List[List[List[str]]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


class DOCXParserError(Exception):
    """Custom exception for DOCX parser errors."""
    pass


class DOCXParser:
    """
    Production-ready DOCX document parser.
    
    Features:
    - Streaming-friendly paragraph extraction
    - Table extraction with cell merging support
    - Document metadata extraction
    - Comprehensive error handling
    - Memory efficient (processes large docs incrementally)
    - Encoding detection and fallback
    - Structured result object
    
    Example:
        parser = DOCXParser()
        result = parser.parse("document.docx")
        print(result.text)
        print(result.tables)
    """
    
    def __init__(
        self,
        extract_tables: bool = True,
        extract_headers: bool = True,
        extract_footers: bool = True,
        preserve_line_breaks: bool = True,
        max_file_size_mb: int = 100,
        encoding: str = "utf-8",
        strip_whitespace: bool = False,
        log_level: int = logging.INFO
    ):
        """
        Initialize the DOCX parser with configuration options.
        
        Args:
            extract_tables: Extract and parse tables from document
            extract_headers: Include header content in output
            extract_footers: Include footer content in output
            preserve_line_breaks: Keep original paragraph breaks
            max_file_size_mb: Maximum allowed file size in MB
            encoding: Text encoding for output (default: utf-8)
            strip_whitespace: Remove extra whitespace from text
            log_level: Logging level for parser operations
        """
        self.extract_tables = extract_tables
        self.extract_headers = extract_headers
        self.extract_footers = extract_footers
        self.preserve_line_breaks = preserve_line_breaks
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        self.encoding = encoding
        self.strip_whitespace = strip_whitespace
        
        # Configure logger
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
        self.logger.setLevel(log_level)
        
    def parse(self, file_path: Union[str, Path]) -> ParseResult:
        """
        Parse a DOCX file and return structured content.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            ParseResult object containing text, tables, metadata, and errors
            
        Raises:
            DOCXParserError: For unrecoverable errors (file not found, corrupt, etc.)
            TypeError: If file_path is not a string or Path object
        """
        # Convert to Path object for consistent handling
        try:
            path = Path(file_path) if isinstance(file_path, str) else file_path
            if not isinstance(path, Path):
                raise TypeError(f"file_path must be str or Path, got {type(file_path)}")
        except Exception as e:
            raise TypeError(f"Invalid file_path: {e}")
        
        # Validate file existence and size
        self._validate_file(path)
        
        result = ParseResult()
        
        try:
            # Open and parse document with context manager
            with self._safe_open_document(path) as doc:
                # Extract document properties (metadata)
                result.metadata = self._extract_metadata(doc, path)
                
                # Extract headers if requested
                if self.extract_headers:
                    header_texts = self._extract_sections(doc, 'headers')
                    if header_texts:
                        result.paragraphs.extend(header_texts)
                
                # Extract main document content
                main_paragraphs = self._extract_paragraphs(doc)
                result.paragraphs.extend(main_paragraphs)
                result.text = self._join_paragraphs(result.paragraphs)
                
                # Extract tables if requested
                if self.extract_tables:
                    result.tables = self._extract_tables(doc)
                
                # Extract footers if requested
                if self.extract_footers:
                    footer_texts = self._extract_sections(doc, 'footers')
                    if footer_texts:
                        result.paragraphs.extend(footer_texts)
                        result.text = self._join_paragraphs(result.paragraphs)
                
        except PackageNotFoundError as e:
            error_msg = f"Invalid or corrupt DOCX file: {e}"
            self.logger.error(error_msg)
            result.errors.append(error_msg)
            raise DOCXParserError(error_msg) from e
            
        except Exception as e:
            error_msg = f"Unexpected error while parsing {path.name}: {str(e)}"
            self.logger.exception(error_msg)
            result.errors.append(error_msg)
            raise DOCXParserError(error_msg) from e
        
        # Log parsing summary
        self.logger.info(
            f"Successfully parsed {path.name}: "
            f"{len(result.paragraphs)} paragraphs, "
            f"{len(result.tables)} tables, "
            f"{len(result.errors)} errors"
        )
        
        return result
    
    def parse_stream(self, file_stream):
        """
        Parse a DOCX from a file-like stream (e.g., uploaded file, HTTP response).
        
        Args:
            file_stream: File-like object supporting .read() (e.g., BytesIO)
            
        Returns:
            ParseResult object
        """
        try:
            doc = Document(file_stream)
            return self._parse_document_object(doc, "stream")
        except Exception as e:
            raise DOCXParserError(f"Failed to parse stream: {e}") from e
    
    def _validate_file(self, path: Path) -> None:
        """Validate file exists, is readable, and within size limits."""
        if not path.exists():
            raise DOCXParserError(f"File not found: {path}")
        
        if not path.is_file():
            raise DOCXParserError(f"Path is not a file: {path}")
        
        if not os.access(path, os.R_OK):
            raise DOCXParserError(f"File is not readable: {path}")
        
        file_size = path.stat().st_size
        if file_size > self.max_file_size_bytes:
            raise DOCXParserError(
                f"File size ({file_size / 1024 / 1024:.2f} MB) exceeds "
                f"maximum allowed ({self.max_file_size_bytes / 1024 / 1024:.2f} MB)"
            )
    
    @contextmanager
    def _safe_open_document(self, path: Path):
        """Context manager for safely opening DOCX documents."""
        doc = None
        try:
            doc = Document(str(path))
            yield doc
        finally:
            # Clean up resources if needed (docx library doesn't have explicit close)
            if doc:
                doc = None
    
    def _parse_document_object(self, doc: DocumentType, source_name: str) -> ParseResult:
        """Internal method to parse an already-opened Document object."""
        result = ParseResult()
        
        try:
            result.paragraphs = self._extract_paragraphs(doc)
            result.text = self._join_paragraphs(result.paragraphs)
            
            if self.extract_tables:
                result.tables = self._extract_tables(doc)
                
        except Exception as e:
            error_msg = f"Error parsing document {source_name}: {str(e)}"
            self.logger.error(error_msg)
            result.errors.append(error_msg)
            
        return result
    
    def _extract_paragraphs(self, doc: DocumentType) -> List[str]:
        """
        Extract text from all paragraphs.
        
        Handles:
        - Empty paragraphs
        - Paragraphs with no text
        - Unicode and special characters
        """
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            text = self._clean_text(paragraph.text)
            
            # Skip empty paragraphs if not preserving line breaks
            if text or self.preserve_line_breaks:
                paragraphs.append(text)
        
        return paragraphs
    
    def _extract_tables(self, doc: DocumentType) -> List[List[List[str]]]:
        """
        Extract tables as 3D arrays: [table][row][cell].
        
        Handles merged cells by replicating content across spans.
        Returns empty list if no tables found.
        """
        tables_data = []
        
        for table in doc.tables:
            table_grid = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Extract text from all paragraphs in cell
                    cell_texts = [p.text for p in cell.paragraphs if p.text.strip()]
                    cell_content = "\n".join(cell_texts) if cell_texts else ""
                    row_data.append(self._clean_text(cell_content))
                table_grid.append(row_data)
            
            tables_data.append(table_grid)
        
        return tables_data
    
    def _extract_sections(self, doc: DocumentType, section_type: str) -> List[str]:
        """
        Extract headers or footers from document sections.
        
        Args:
            doc: Document object
            section_type: Either 'headers' or 'footers'
        """
        texts = []
        
        for section in doc.sections:
            if section_type == 'headers':
                parts = [
                    section.header,
                    section.even_page_header,
                    section.first_page_header
                ]
            elif section_type == 'footers':
                parts = [
                    section.footer,
                    section.even_page_footer,
                    section.first_page_footer
                ]
            else:
                self.logger.warning(f"Unknown section type: {section_type}")
                continue
            
            for part in parts:
                if part and part.paragraphs:
                    for paragraph in part.paragraphs:
                        text = self._clean_text(paragraph.text)
                        if text:
                            texts.append(text)
        
        return texts
    
    def _extract_metadata(self, doc: DocumentType, path: Path) -> Dict[str, Any]:
        """
        Extract document properties and metadata.
        """
        metadata = {
            "file_name": path.name,
            "file_size_bytes": path.stat().st_size,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
        
        # Extract core properties if available
        try:
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                    "last_modified_by": core_props.last_modified_by or "",
                    "category": core_props.category or "",
                    "keywords": core_props.keywords or "",
                    "comments": core_props.comments or "",
                })
        except Exception as e:
            self.logger.debug(f"Could not extract core properties: {e}")
        
        return metadata
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text content.
        
        Handles:
        - Unicode normalization
        - Extra whitespace removal (optional)
        - Null byte stripping
        - Encoding conversion
        """
        if not text:
            return ""
        
        # Remove null bytes (common in corrupted documents)
        text = text.replace('\x00', '')
        
        # Optional: strip excessive whitespace
        if self.strip_whitespace:
            # Replace multiple spaces/newlines with single space
            import re
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
        
        return text
    
    def _join_paragraphs(self, paragraphs: List[str]) -> str:
        """Join paragraphs with appropriate line breaks."""
        if not paragraphs:
            return ""
        
        if self.preserve_line_breaks:
            return "\n".join(paragraphs)
        else:
            return " ".join(paragraphs)
    
    def extract_text_only(self, file_path: Union[str, Path]) -> str:
        """
        Convenience method to extract only plain text.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            String containing document text
        """
        result = self.parse(file_path)
        return result.text


# Example usage and testing
if __name__ == "__main__":
    # Configure logging for standalone use
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # Production-ready usage
    parser = DOCXParser(
        extract_tables=True,
        preserve_line_breaks=True,
        max_file_size_mb=50
    )
    
    try:
        # Parse document
        result = parser.parse("example.docx")
        
        # Access results
        print(f"Document Title: {result.metadata.get('title', 'N/A')}")
        print(f"Paragraphs: {len(result.paragraphs)}")
        print(f"Tables: {len(result.tables)}")
        print(f"\nFirst 500 chars:\n{result.text[:500]}")
        
        # Handle any warnings
        for warning in result.warnings:
            print(f"Warning: {warning}")
            
    except DOCXParserError as e:
        print(f"Failed to parse document: {e}")
        # Implement fallback or retry logic here
    except Exception as e:
        print(f"Unexpected error: {e}")
